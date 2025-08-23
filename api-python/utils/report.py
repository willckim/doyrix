# report.py
import logging
import re
from typing import Dict, Any, Tuple, List
from jinja2 import Template

log = logging.getLogger("doyrix")

# robust import: works if ai.py is in root or in utils/
try:
    from utils.ai import summarize_section
except ImportError:
    from ai import summarize_section

# --- numeric detection (for right-aligning numbers) --------------------------------
_NUM_RE = re.compile(r"""
    ^\s*
    \(?        # optional opening paren
    \$?        # optional dollar
    -?         # optional minus
    \d[\d,]*   # digits with thousands sep
    (?:\.\d+)? # optional decimals
    %?         # optional percent
    \)?        # optional closing paren
    \s*$
""", re.X)

def is_num(cell: Any) -> bool:
    s = (str(cell) if cell is not None else "").strip()
    if s in {"", "-", "—", "–"}:
        return True
    return bool(_NUM_RE.match(s))

def is_key_label(label: Any) -> bool:
    """Decide if a row label is a 'key line' to surface above the full table."""
    s = (str(label) if label is not None else "").strip().lower()
    if not s:
        return False
    return (
        s.startswith("total")
        or s.startswith("net")
        or "cash and cash equivalents" in s
        or "revenue" in s
        or "operating income" in s
        or "operating loss" in s
        or "gross profit" in s
        or "gross margin" in s
    )

REPORT_TEMPLATE = Template("""
<article style='font-family:ui-sans-serif;max-width:1000px;margin:auto;padding:24px;'>
  <style>
  /* Print-friendly tweaks */
  @media print {
    .page-break { page-break-before: always; }
    table { page-break-inside: auto; }
    tr, td, th { page-break-inside: avoid; page-break-after: auto; }
    h1, h2, h3 { page-break-after: avoid; }
  }
  /* On-screen readability */
  code { background:#f1f5f9; padding:0 .25rem; border-radius:4px; }
  table { border-collapse: collapse; width:100%; }
  th, td { border-top:1px solid #e5e7eb; padding:6px 8px; }
  thead th { border-bottom:1px solid #e2e8f0; }

  /* Sticky TOC + anchor offset */
  :root { --toc-h: 44px; }
  nav.report-toc {
    position: sticky;
    top: 0;
    z-index: 50;
    background: #fff;
  }
  section { scroll-margin-top: calc(var(--toc-h) + 8px); }

  /* In print, don't stick the nav */
  @media print {
    nav.report-toc { position: static !important; }
  }
</style>

  <header style="margin-bottom:16px;">
    <h1 style="margin:0;">Analyst Report</h1>
    <p style="color:#475569;margin:.5rem 0 0 0;">
      <strong>Pages:</strong> {{ meta.pages }} · <strong>Detected:</strong> {{ meta.doc_type }}
      {% if meta.company or meta.ticker or meta.filing_date %} ·
        {% if meta.company %}<strong>{{ meta.company }}</strong>{% endif %}
        {% if meta.ticker %} ({{ meta.ticker }}){% endif %}
        {% if meta.filing_date %} — {{ meta.filing_date }}{% endif %}
      {% endif %}
    </p>
  </header>

  <nav class="report-toc" style="padding:8px 0;border-top:1px solid #e5e7eb;border-bottom:1px solid #e5e7eb;margin-bottom:16px;">
    <a href="#overview" style="margin-right:12px;">Overview</a>
    <a href="#mdna" style="margin-right:12px;">MD&A</a>
    <a href="#financials" style="margin-right:12px;">Financials</a>
    <a href="#non-gaap" style="margin-right:12px;">Non-GAAP</a>
    <a href="#capital" style="margin-right:12px;">Capital Structure</a>
    <a href="#risks" style="margin-right:12px;">Risks</a>
    <a href="#market-risk" style="margin-right:12px;">Market Risk</a>
    <a href="#controls" style="margin-right:12px;">Controls/Auditor</a>
    <a href="#legal" style="margin-right:12px;">Legal/Contingencies</a>
    <a href="#appendix">Appendix</a>
  </nav>

  <!-- Overview -->
  <section id="overview" style="margin:20px 0;">
    <h2>Overview</h2>
    <p class="muted" style="color:#334155">Auto-generated highlights with page references like <code>[p12]</code>.</p>
    {% if kpis %}
      <h3>Key Metrics</h3>
      <table style="font-size:14px;margin-bottom:12px;">
        <thead>
          <tr>
            <th style="text-align:left;">Metric</th>
            <th style="text-align:right;">Value</th>
            <th style="text-align:right;">YoY</th>
            <th style="text-align:right;">Pages</th>
          </tr>
        </thead>
        <tbody>
          {% for k in kpis %}
          <tr>
            <td>{{ k.name.replace('_',' ').title() }}</td>
            <td style="text-align:right;">{{ k.value }}</td>
            <td style="text-align:right;">{% if k.yoy is not none %}{{ (k.yoy * 100) | round(1) }}%{% else %}—{% endif %}</td>
            <td style="text-align:right;">{% if k.pages %}[p{{ k.pages[0] }}]{% else %}—{% endif %}</td>
          </tr>
          {% endfor %}
        </tbody>
      </table>
    {% endif %}

    {% if segments %}
      <h3>Segments & Geography</h3>
      <table style="font-size:14px;margin-bottom:12px;">
        <thead>
          <tr>
            <th style="text-align:left;">Segment/Region</th>
            <th style="text-align:right;">Revenue</th>
            <th style="text-align:right;">Gross Margin</th>
            <th style="text-align:right;">Pages</th>
          </tr>
        </thead>
        <tbody>
          {% for seg in segments %}
          <tr>
            <td>{{ seg.name }}</td>
            <td style="text-align:right;">{{ seg.rev }}</td>
            <td style="text-align:right;">{{ seg.gm }}</td>
            <td style="text-align:right;">{% if seg.pages %}[p{{ seg.pages[0] }}]{% endif %}</td>
          </tr>
          {% endfor %}
        </tbody>
      </table>
    {% endif %}
  </section>

  <div class="page-break"></div>

  <!-- MD&A -->
  <section id="mdna" style="margin:24px 0%;">
    <h2>Management’s Discussion & Analysis (MD&A)</h2>
    {% if mdna_summary %}
      <ul style="line-height:1.5;padding-left:1.1rem;">
        {% for line in mdna_summary.splitlines() if line.strip() %}
          <li>{{ line | e }}</li>
        {% endfor %}
      </ul>
    {% else %}
      <p style="color:#64748b;">(No MD&A summary extracted)</p>
    {% endif %}
  </section>

  <div class="page-break"></div>

  <!-- Financial Statements -->
  <section id="financials" style="margin:24px 0;">
    <h2>Financial Statements</h2>

    <style>
      /* Table readability upgrades */
      .fin-wrap { overflow:auto; border:1px solid #e5e7eb; border-radius:8px; }
      .fin-table { border-collapse: collapse; width:100%; font-size:13px; }
      .fin-table thead th {
        position: sticky; top: 0; background: #fff;
        border-bottom:1px solid #e2e8f0;
      }
      .fin-table th, .fin-table td { padding:6px 8px; border-top:1px solid #e5e7eb; }
      .fin-table tbody tr:nth-child(odd) { background:#fafafa; }
      .fin-table td.num { text-align:right; font-variant-numeric: tabular-nums; white-space:nowrap; }
      .fin-table td.label { max-width: 40ch; }
      details.fin { margin:8px 0 20px; }
      details.fin > summary { cursor:pointer; color:#334155; }
    </style>

    {% if financials %}
      {% for table in financials %}
        <h3 style="margin-top:12px;">
          {{ table.title or "Table" }}{% if table.pages %} [p{{ table.pages[0] }}]{% endif %}
        </h3>

        {# ---- Optional compact "key lines" view first ---- #}
        {% set ns = namespace(rows=[]) %}
        {% for r in table.rows %}
          {% set label = (r[0] if r and r[0] is string else "") %}
          {% if label and is_key_label(label) %}
            {% set ns.rows = ns.rows + [r] %}
          {% endif %}
        {% endfor %}

        {% if ns.rows %}
          <div class="fin-wrap">
            <table class="fin-table">
              {% if table.header %}
                <thead>
                  <tr>
                    {% for h in table.header %}
                      <th style="text-align:left;">{{ h }}</th>
                    {% endfor %}
                  </tr>
                </thead>
              {% endif %}
              <tbody>
                {% for row in ns.rows %}
                  <tr>
                    {% for cell in row %}
                      {% set numeric = is_num(cell) %}
                      <td class="{{ 'num' if numeric else 'label' if loop.index0 == 0 else '' }}">{{ cell }}</td>
                    {% endfor %}
                  </tr>
                {% endfor %}
              </tbody>
            </table>
          </div>
          <p style="color:#64748b; margin:6px 0 0 0;">Key lines shown. Full table below.</p>
        {% endif %}

        {# ---- Full table collapsible ---- #}
        <details class="fin" {% if not ns.rows %}open{% endif %}>
          <summary>Show full {{ table.title or "table" }}</summary>
          <div class="fin-wrap" style="margin-top:8px;">
            <table class="fin-table">
              {% if table.header %}
                <thead>
                  <tr>
                    {% for h in table.header %}
                      <th style="text-align:left;">{{ h }}</th>
                    {% endfor %}
                  </tr>
                </thead>
              {% endif %}
              <tbody>
                {% for row in table.rows %}
                  <tr>
                    {% for cell in row %}
                      {% set numeric = is_num(cell) %}
                      <td class="{{ 'num' if numeric else 'label' if loop.index0 == 0 else '' }}">{{ cell }}</td>
                    {% endfor %}
                  </tr>
                {% endfor %}
              </tbody>
            </table>
          </div>
        </details>

        {% if loop.index % 2 == 0 %}
          <div class="page-break"></div>
        {% endif %}
      {% endfor %}
    {% else %}
      <p style="color:#64748b;">(No tables parsed yet — add table extraction to the parser)</p>
    {% endif %}
  </section>

  <div class="page-break"></div>

  <!-- Non-GAAP -->
  <section id="non-gaap" style="margin:24px 0;">
    <h2>Non-GAAP Reconciliations</h2>
    {% if non_gaap %}
      {% for rec in non_gaap %}
        <h3>{{ rec.metric }} — {{ rec.period }}{% if rec.pages %} [p{{ rec.pages[0] }}]{% endif %}</h3>
        <ul style="line-height:1.5;padding-left:1.1rem;">
          {% for r in rec.recon %}
            <li>{{ r.label }}: {{ r.value }}</li>
          {% endfor %}
        </ul>
      {% endfor %}
    {% else %}
      <p style="color:#64748b;">(No non-GAAP reconciliations extracted)</p>
    {% endif %}
  </section>

  <div class="page-break"></div>

  <!-- Capital Structure -->
  <section id="capital" style="margin:24px 0;">
    <h2>Capital Structure</h2>
    {% if capital %}
      <ul style="line-height:1.6;padding-left:1.1rem;">
        {% if capital.cash is defined %}<li>Cash & equivalents: {{ capital.cash }}</li>{% endif %}
        {% if capital.total_debt is defined %}<li>Total debt: {{ capital.total_debt }}</li>{% endif %}
        {% if capital.net_cash is defined %}<li>Net cash (debt): {{ capital.net_cash }}</li>{% endif %}
        {% if capital.facilities is defined %}<li>Facilities (undrawn): {{ capital.facilities }}</li>{% endif %}
      </ul>
      {% if capital.instruments %}
        <h3>Debt Instruments</h3>
        <div style="overflow:auto;border:1px solid #e5e7eb;border-radius:8px;">
          <table style="font-size:13px;">
            <thead>
              <tr>
                <th style="text-align:left;">Name</th>
                <th style="text-align:right;">Coupon</th>
                <th style="text-align:left;">Currency</th>
                <th style="text-align:left;">Maturity</th>
                <th style="text-align:right;">Amount</th>
                <th style="text-align:left;">Pages</th>
              </tr>
            </thead>
            <tbody>
              {% for inst in capital.instruments %}
              <tr>
                <td>{{ inst.name }}</td>
                <td style="text-align:right;">{{ inst.coupon }}</td>
                <td>{{ inst.currency }}</td>
                <td>{{ inst.maturity }}</td>
                <td style="text-align:right;">{{ inst.amount }}</td>
                <td>{% if inst.pages %}[p{{ inst.pages[0] }}]{% endif %}</td>
              </tr>
              {% endfor %}
            </tbody>
          </table>
        </div>
      {% endif %}
      {% if capital.maturity_ladder %}
        <h3>Maturity Ladder</h3>
        <ul style="line-height:1.5;padding-left:1.1rem;">
          {% for step in capital.maturity_ladder %}
            <li>{{ step.year }} — {{ step.amount }}</li>
          {% endfor %}
        </ul>
      {% endif %}
    {% else %}
      <p style="color:#64748b;">(No capital structure parsed yet)</p>
    {% endif %}
  </section>

  <div class="page-break"></div>

  <!-- Risks -->
  <section id="risks" style="margin:24px 0;">
    <h2>Top Risks</h2>
    {% if risks %}
      <ul style="line-height:1.5;padding-left:1.1rem;">
        {% for r in risks %}
          <li>{{ r.text | e }} {% if r.page %}[p{{ r.page }}]{% endif %}</li>
        {% endfor %}
      </ul>
    {% else %}
      <p style="color:#64748b;">(No risk factors summarized)</p>
    {% endif %}
  </section>

  <div class="page-break"></div>

  <!-- Market Risk -->
  <section id="market-risk" style="margin:24px 0;">
    <h2>Market Risk (Item 7A)</h2>
    {% if market_risk %}
  {% for k,v in market_risk.items() %}
  <h3 style="margin-top:8px;">{{ k.replace('_',' ').title() }}</h3>
  {% if v is string %}
    <p>{{ v }}</p>
  {% elif v is sequence %}
    <ul style="line-height:1.5;padding-left:1.1rem;">
      {% for line in v if line %}
        <li>{{ line }}</li>
      {% endfor %}
    </ul>
  {% elif v is mapping %}
    <ul style="line-height:1.5;padding-left:1.1rem;">
      {% for kk, vv in v.items() %}
        <li><strong>{{ kk.replace('_',' ').title() }}:</strong> {{ vv }}</li>
      {% endfor %}
    </ul>
  {% else %}
    <pre style="background:#f8fafc;border:1px solid #e5e7eb;border-radius:8px;padding:8px;white-space:pre-wrap;">{{ v | tojson }}</pre>
  {% endif %}
{% endfor %}
{% else %}
  <p style="color:#64748b;">(No market risk details extracted)</p>
{% endif %}
  </section>

  <div class="page-break"></div>

  <!-- Controls/Auditor -->
  <section id="controls" style="margin:24px 0;">
    <h2>Controls & Auditor</h2>
    {% if auditor %}
      <ul style="line-height:1.5;padding-left:1.1rem;">
        <li>Opinion: {{ auditor.opinion or "—" }}</li>
        {% if auditor.get('cam') %}<li>Critical Audit Matters: {{ auditor.get('cam') | length }}</li>{% endif %}
        {% if auditor.pages %}<li>Pages: [p{{ auditor.pages[0] }}]</li>{% endif %}
      </ul>
    {% else %}
      <p style="color:#64748b;">(No auditor/control details extracted)</p>
    {% endif %}
  </section>

  <div class="page-break"></div>

    <!-- Legal/Contingencies -->
  <section id="legal" style="margin:24px 0;">
    <h2>Legal, Contingencies & Subsequent Events</h2>
    {% set legal_items = legal.get('items', []) %}
    {% if legal_items %}
      <ul style="line-height:1.5;padding-left:1.1rem;">
        {% for l in legal_items %}
          <li>{{ l.title }} — {{ l.summary }} {% if l.pages %}[p{{ l.pages[0] }}]{% endif %}</li>
        {% endfor %}
      </ul>
    {% else %}
      <p style="color:#64748b;">(No legal/contingency items extracted)</p>
    {% endif %}
  </section>

  <div class="page-break"></div>

  <!-- Appendix -->
  <section id="appendix" style="margin:24px 0;">
    <h2>Appendix: Section Summaries & Citations</h2>
    {% for s in sections %}
      <h3 style="margin:.5rem 0;">{{ s.title }}</h3>
      <p style="color:#64748b;margin:.25rem 0;">p.{{ s.start_page }}–{{ s.end_page }}</p>
      {% if s.summary %}
        <ul style="line-height:1.35;padding-left:1.1rem;">
          {% for line in s.summary.splitlines() if line.strip() %}
            <li>{{ line | e }}</li>
          {% endfor %}
        </ul>
      {% else %}
        <p style="color:#94a3b8;">(No summary available)</p>
      {% endif %}
    {% endfor %}

    {% if citations %}
      <h3 style="margin-top:1rem;">Sample Citations</h3>
      <ol style="padding-left:1.2rem; color:#334155;">
        {% for c in citations[:12] %}
          <li>p.{{ c.page }} — {{ c.snippet | e }}</li>
        {% endfor %}
      </ol>
    {% endif %}
  </section>
</article>
""")

def _first_chunks_with_pages(
    content: List[dict],
    chars_per_chunk: int = 1600,
    max_chunks: int | None = None
):
    pages_list = [it.get("page") for it in content if "page" in it]
    if max_chunks is None:
        span = (max(pages_list) - min(pages_list) + 1) if pages_list else 1
        max_chunks = 3 if span <= 6 else (5 if span <= 12 else 6)

    chunks, pages = [], []
    buf, buf_page = "", None
    for item in content:
        text = (item.get("full") or item.get("snippet") or "").strip()
        if not text:
            continue
        if buf == "":
            buf_page = item.get("page", buf_page)
        if len(buf) + len(text) <= chars_per_chunk:
            buf += ("\n" if buf else "") + text
        else:
            if buf:
                chunks.append(buf); pages.append(buf_page or item.get("page", 1))
            buf = text[:chars_per_chunk]; buf_page = item.get("page", buf_page)
            if len(chunks) >= max_chunks:
                break
    if buf and len(chunks) < max_chunks:
        chunks.append(buf); pages.append(buf_page or (pages_list[0] if pages_list else 1))
    return chunks, pages

def _normalize_quotes(s: str) -> str:
    if not isinstance(s, str):
        return ""
    # normalize curly quotes for robust matching
    return s.replace("’", "'").replace("‘", "'").replace("“", '"').replace("”", '"')

def _summarize_sections(parsed: Dict[str, Any]) -> List[Dict[str, Any]]:
    out = []
    for s in parsed.get("sections", []):
        content_items = s.get("content", [])
        chunks, pages = _first_chunks_with_pages(content_items)
        try:
            log.info("Summarize: %r pages=%s-%s items=%d chunks=%d",
                     s.get("title"), s.get("start_page"), s.get("end_page"),
                     len(content_items), len(chunks))
        except Exception:
            pass
        summary = ""
        if chunks:
            summary = summarize_section(s.get("title", "Section"), chunks, pages) or ""
        if not summary.strip():
            if content_items:
                raw = (content_items[0].get("snippet") or "").strip()
                lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
                summary = "\n".join(f"- {ln}" for ln in lines[:3]) if lines else "(no text extracted)"
            else:
                summary = "(no content in section)"
        out.append({**s, "summary": summary})
    return out

def build_analyst_report(parsed: Dict[str, Any]) -> Tuple[str, dict]:
    sections = _summarize_sections(parsed)
    meta = {
        "pages": parsed.get("doc_meta", {}).get("pages", 0),
        "doc_type": parsed.get("doc_meta", {}).get("doc_type", "unknown"),
        "anchors_found": parsed.get("doc_meta", {}).get("anchors_found"),
        "company": (parsed.get("company") or parsed.get("meta", {}).get("company")),
        "ticker": (parsed.get("ticker") or parsed.get("meta", {}).get("ticker")),
        "filing_date": (parsed.get("filing_date") or parsed.get("meta", {}).get("filing_date")),
    }

    # Pull structured “derived” data if the parser populates it
    derived = parsed.get("derived", {})
    kpis = derived.get("kpis", [])
    segments = derived.get("segments", [])
    financials = derived.get("financials", [])          # list of {title, header, rows, pages}
    non_gaap = derived.get("non_gaap", [])              # list of {metric, period, recon:[{label,value}], pages}
    capital = parsed.get("capital_structure") or derived.get("capital_structure") or {}
    risks = derived.get("risks") or []                  # list of {text, page}
    market_risk = derived.get("market_risk") or {}
    auditor = derived.get("auditor") or {}
    legal = derived.get("legal") or {}

    # Try to find MD&A section to produce a focused summary block
    mdna_summary = ""
    for s in sections:
        title_norm = _normalize_quotes(s.get("title") or "").lower()
        if "management's discussion" in title_norm:
            mdna_summary = s.get("summary") or ""
            break

    html = REPORT_TEMPLATE.render(
        meta=meta,
        sections=sections,
        citations=parsed.get("citations", []),
        kpis=kpis,
        segments=segments,
        financials=financials,
        non_gaap=non_gaap,
        capital=capital,
        risks=risks,
        market_risk=market_risk,
        auditor=auditor,
        legal=legal,
        mdna_summary=mdna_summary,
        is_num=is_num,
        is_key_label=is_key_label,  # <<< pass helper into template
    )
    metadata = {
        "citations_count": len(parsed.get("citations", [])),
        "sections": [{"title": s["title"], "start_page": s.get("start_page"), "end_page": s.get("end_page")} for s in sections],
        "kpi_count": len(kpis),
        "financial_tables": len(financials),
        "non_gaap_items": len(non_gaap),
        "risk_count": len(risks),
    }
    return html, metadata

# --- EXPORT (PDF/DOCX) using same strategy as one_pager ---
from pathlib import Path

def export_report(parsed: Dict[str, Any], out_path: str | Path, fmt: str = "pdf") -> Path:
    """
    Render Analyst Report to PDF/DOCX.
    PDF: WeasyPrint -> fallback Playwright (Chromium)
    DOCX: python-docx + bs4 (simple mapping)
    """
    out_path = Path(out_path)
    fmt = fmt.lower()
    html, _meta = build_analyst_report(parsed)

    if fmt == "pdf":
        # Try WeasyPrint first
        try:
            from weasyprint import HTML  # may require native deps
            HTML(string=html, base_url=".").write_pdf(str(out_path))
            return out_path
        except Exception as e_wp:
            # Fallback to Playwright
            try:
                from playwright.sync_api import sync_playwright
                with sync_playwright() as p:
                    browser = p.chromium.launch()
                    page = browser.new_page()
                    page.set_content(html, wait_until="load")
                    page.pdf(
                        path=str(out_path),
                        format="Letter",
                        print_background=True,
                        margin={"top": "0.5in", "right": "0.5in", "bottom": "0.5in", "left": "0.5in"},
                    )
                    browser.close()
                return out_path
            except Exception as e_pw:
                html_fallback = out_path.with_suffix(".html")
                html_fallback.write_text(html, encoding="utf-8")
                raise RuntimeError(
                    "PDF render failed with WeasyPrint and Playwright. "
                    f"Wrote HTML to {html_fallback}. "
                    f"WeasyPrint error: {e_wp}; Playwright error: {e_pw}"
                )

    if fmt == "docx":
        try:
            from docx import Document
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            doc = Document()
            for node in soup.find_all(["h1", "h2", "h3", "p", "li"]):
                text = node.get_text(strip=True)
                if not text:
                    continue
                tag = node.name.lower()
                if tag == "h1":
                    doc.add_heading(text, level=0)
                elif tag == "h2":
                    doc.add_heading(text, level=1)
                elif tag == "h3":
                    doc.add_heading(text, level=2)
                elif tag == "li":
                    doc.add_paragraph(f"• {text}")
                else:
                    doc.add_paragraph(text)
            doc.save(str(out_path))
            return out_path
        except Exception as e:
            html_fallback = out_path.with_suffix(".html")
            html_fallback.write_text(html, encoding="utf-8")
            raise RuntimeError(
                f"DOCX export failed ({e}). Wrote HTML to {html_fallback}. "
                f"Install `python-docx` and `beautifulsoup4`."
            )

    raise ValueError("fmt must be 'pdf' or 'docx'")
